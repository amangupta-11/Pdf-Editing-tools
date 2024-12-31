import os
import zipfile
from flask import Flask, request, send_file, render_template, jsonify
from io import BytesIO
import PyPDF2
import fitz  # PyMuPDF for PDF operations
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib.utils import ImageReader
import tempfile
from pdf2docx import Converter
from PIL import Image
import img2pdf
from PyPDF2 import PdfReader, PdfWriter
import pikepdf
from docx import Document
import pythoncom
import psutil
from docx2pdf import convert

import sys

# Conditionally import pythoncom only if on Windows
if sys.platform == "win32":
    import pythoncom


app = Flask(__name__)
# Index route
@app.route('/')
def index():
    return render_template('index.html')

# ------------------------- Merge PDFs -------------------------
@app.route('/file_selection/merge_pdfs', methods=['GET', 'POST'], endpoint='file_selection_merge_pdfs')
def file_selection_merge_pdfs():
    if request.method == 'POST':       
            files = [request.files[key] for key in request.files]
            return merge_pdfs(files)
    return render_template('file-selection.html', tool='merge_pdfs')

def merge_pdfs(files):
    if len(files) < 2:
        return "Please upload at least two PDF files to merge.", 400

    pdf_writer = PyPDF2.PdfWriter()
    merged_file_path = os.path.join('output', 'merged_output.pdf')

    # Process each uploaded PDF file
    for file in files:
        if file.filename == '':
            continue  # Skip empty selections if any

        temp_file_path = os.path.join('uploads', file.filename)
        file.save(temp_file_path)  # Save each file temporarily

        # Open and append each PDF to the writer
        with open(temp_file_path, "rb") as pdf_file:
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            for page in pdf_reader.pages:
                pdf_writer.add_page(page)

        # Optionally remove the temporary file after processing
        os.remove(temp_file_path)

    # Save the merged PDF file
    with open(merged_file_path, "wb") as output_file:
        pdf_writer.write(output_file)

    # Send the merged file to the user
    return send_file(merged_file_path, as_attachment=True, download_name='merged_output.pdf')


# ------------------------- JPG to PDF -------------------------
@app.route('/file_selection/jpg_to_pdf', methods=['GET'],endpoint='file_selection_jpg_to_pdf')
def file_selection_jpg_to_pdf():
    return render_template('file-selection.html',tool='jpg_to_pdf')

@app.route('/convert_to_pdf', methods=['POST'])
def convert_to_pdf():
    files = request.files.getlist('files')
    if not files:
        return "No files uploaded", 400
    try:
        # Create a BytesIO object to store the resulting PDF
        pdf_output = BytesIO()

        # Convert the images to PDF and save in pdf_output
        pdf_output.write(img2pdf.convert([file.stream for file in files]))
        pdf_output.seek(0)

        # Send the resulting PDF to the user
        return send_file(
            pdf_output,
            as_attachment=True,
            download_name="output.pdf",
            mimetype="application/pdf"
        )
    except Exception as e:
        return f"Error processing files: {e}", 500


# ------------------------- PDF to JPG -------------------------
@app.route('/file_selection/pdf_to_jpg', methods=['GET'],endpoint='file_selection_pdf_to_jpg')
def file_selection_pdf_to_jpg():
    return render_template('file-selection.html',tool='pdf_to_jpg')

@app.route('/convert_to_jpg', methods=['POST'])
def convert_to_jpg():
    pdf_file = request.files.get('file')
    if not pdf_file:
        return "No PDF file uploaded", 400

    try:
        # Load the PDF using PyMuPDF from a BytesIO stream
        pdf_stream = BytesIO(pdf_file.read())
        pdf_document = fitz.open(stream=pdf_stream, filetype="pdf")
        image_outputs = []

        # Convert each page to an image
        for page_number in range(len(pdf_document)):
            page = pdf_document[page_number]
            pix = page.get_pixmap(dpi=300)  # Render the page as an image (300 DPI for better quality)
            img_stream = BytesIO(pix.tobytes(output="jpeg"))  # Convert the image to a BytesIO object
            image_outputs.append((f"page_{page_number + 1}.jpg", img_stream))

        # Return the images as a zip file
        zip_output = BytesIO()
        with zipfile.ZipFile(zip_output, 'w') as zip_file:
            for filename, img_stream in image_outputs:
                zip_file.writestr(filename, img_stream.getvalue())
        zip_output.seek(0)

        return send_file(
            zip_output,
            as_attachment=True,
            download_name="pdf_to_jpg.zip",
            mimetype="application/zip"
        )

    except Exception as e:
        return f"Error processing PDF: {e}", 500

# ------------------------- Add Page Numbers -------------------------
@app.route('/file_selection/add_page_number', methods=['GET'])
def file_selection_add_page_number():
    # Renders the page for adding page numbers
    return render_template('file-selection.html', tool='add_page_numbers')

@app.route('/add_page_numbers', methods=['POST'])
def add_page_numbers():
    # Handles adding page numbers to the uploaded PDF
    pdf_file = request.files.get('file')
    position = request.form.get('position')  # Get the position from the form

    if not pdf_file:
        return "No PDF file uploaded", 400
    if not position:
        return "No position selected", 400

    try:
        # Read the uploaded PDF file into a stream
        pdf_stream = BytesIO(pdf_file.read())
        pdf_document = fitz.open(stream=pdf_stream, filetype="pdf")
        output_stream = BytesIO()

        # Map user-specified positions to coordinates
        positions = {
            "bottom-left": lambda page: (50, page.rect.height - 30),
            "bottom-right": lambda page: (page.rect.width - 100, page.rect.height - 30),
            "top-left": lambda page: (50, 50),
            "top-right": lambda page: (page.rect.width - 100, 50),
        }

        # Validate the selected position
        if position not in positions:
            return "Invalid position selected", 400

        # Add page numbers to each page at the selected position
        for page_number in range(len(pdf_document)):
            page = pdf_document[page_number]
            text = f"{page_number + 1}"
            font_size = 12
            text_position = positions[position](page)
            page.insert_text(
                text_position,
                text,
                fontsize=font_size,
                color=(0, 0, 0)  # Black color
            )

        # Save the modified PDF to an output stream
        pdf_document.save(output_stream)
        pdf_document.close()
        output_stream.seek(0)

        # Return the processed file as a downloadable attachment
        return send_file(output_stream, as_attachment=True, download_name="page_numbered.pdf", mimetype="application/pdf")

    except Exception as e:
        # Handle and display errors during processing
        return f"Error processing PDF: {e}", 500

# ------------------------- Rotate PDF -------------------------
@app.route('/file_selection/<tool>', methods=['GET'],endpoint='file_selection_rotate_pdf')
def file_selection_rotate_pdf(tool):
    return render_template('file-selection.html', tool=tool)

@app.route('/rotate_pdf', methods=['POST'])
def rotate_pdf():
    file = request.files['file']
    direction = request.form['direction']

    pdf_reader = PdfReader(file)
    pdf_writer = PdfWriter()

    for page in pdf_reader.pages:
        if direction == 'left':
            page.rotate(90)  # Rotate 90° counterclockwise
        elif direction == 'right':
            page.rotate(-90)  # Rotate 90° clockwise
        pdf_writer.add_page(page)

    output = BytesIO()
    pdf_writer.write(output)
    output.seek(0)
    return send_file(output, as_attachment=True, download_name='rotated_output.pdf', mimetype='application/pdf')

# ------------------------- Split PDF -------------------------
@app.route('/file_selection/split_pdf', methods=['GET'],endpoint='file_selection_split_pdf')
def file_selection_split_pdf():
    return render_template('file-selection.html',tool='split_pdf')

def split_pdf(input_pdf_stream, page_range):
    pdf_reader = PdfReader(input_pdf_stream)
    pdf_writer = PdfWriter()
    output_pdfs = []

    # Parse the page range input
    page_numbers = []
    for part in page_range.split(','):
        if '-' in part:
            start, end = part.split('-')
            page_numbers.extend(range(int(start), int(end) + 1))
        else:
            page_numbers.append(int(part))

    # Adjust for 0-based indexing in PyPDF2
    page_numbers = [p - 1 for p in page_numbers]

    # Split the PDF based on the page range
    for page_num in page_numbers:
        pdf_writer = PdfWriter()
        pdf_writer.add_page(pdf_reader.pages[page_num])

        # Save each page as a separate PDF
        output_pdf_stream = BytesIO()
        pdf_writer.write(output_pdf_stream)
        output_pdf_stream.seek(0)
        output_pdfs.append(output_pdf_stream)

    return output_pdfs

@app.route('/split_pdf', methods=['POST'])
def split_pdf_file():
    pdf_file = request.files.get('file')
    if not pdf_file:
        return "No PDF file uploaded", 400

    page_range = request.form.get('page_range')
    if not page_range:
        return "No page range provided", 400

    try:
        # Read the uploaded PDF file
        pdf_stream = BytesIO(pdf_file.read())

        # Split the PDF based on the provided page range
        output_pdfs = split_pdf(pdf_stream, page_range)

        # Create a zip file containing all the split PDFs
        zip_buffer = BytesIO()
        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
            for i, output_pdf in enumerate(output_pdfs):
                zip_file.writestr(f"page_{i+1}.pdf", output_pdf.read())

        zip_buffer.seek(0)
        return send_file(zip_buffer, as_attachment=True, download_name="split_pdfs.zip", mimetype="application/zip")

    except Exception as e:
        return f"Error splitting PDF: {e}", 500

# ------------------------- Protect PDF (Encrypt) -------------------------

# Ensure these directories exist
os.makedirs('uploads', exist_ok=True)
os.makedirs('protected', exist_ok=True)

@app.route('/file_selection/protect_pdf', methods=['POST'])
def protect_pdf():
    file = request.files['file']
    password = request.form['password']
    confirm_password = request.form['confirm_password']
    
    if password != confirm_password:
        return "Passwords do not match", 400

    # Save the uploaded file to a temporary location
    temp_file_path = os.path.join('uploads', file.filename)
    file.save(temp_file_path)

    # Protect the PDF
    protected_file_path = os.path.join('protected', file.filename)
    try:
        with open(temp_file_path, "rb") as pdf_file:
            reader = PyPDF2.PdfReader(pdf_file)
            writer = PyPDF2.PdfWriter()
            
            # Add all pages to the writer
            for page in range(len(reader.pages)):
                writer.add_page(reader.pages[page])
            
            # Encrypt the PDF with the provided password
            writer.encrypt(user_password=password, owner_password=password, use_128bit=True)

            # Save the protected PDF
            with open(protected_file_path, "wb") as protected_file:
                writer.write(protected_file)
    except Exception as e:
        return str(e), 500
    finally:
        # Optionally, remove the original unprotected file
        os.remove(temp_file_path)

    # Send the protected PDF file for download
    return send_file(protected_file_path, as_attachment=True)

@app.route('/file_selection/<tool>', methods=['GET'])
def file_selection(tool):
    return render_template('file-selection.html', tool=tool)


# ------------------------- Unlock PDF  -------------------------

@app.route('/file_selection/unlock_pdf', methods=['GET'],endpoint='file_selection_unlock_pdf')
def file_selection_unlock_pdf():
    return render_template('file-selection.html',tool='unlock_pdf')

@app.route('/check_pdf_protection', methods=['POST'])
def check_pdf_protection():
    pdf_file = request.files.get('file')
    if not pdf_file:
        return jsonify({"error": "No PDF file uploaded"}), 400

    try:
        # Read the uploaded PDF file into a stream
        pdf_stream = BytesIO(pdf_file.read())

        # Attempt to open the PDF without a password
        try:
            with pikepdf.open(pdf_stream):
                # If we can open it without a password, it's not protected
                return jsonify({"protected": False})
        except pikepdf.PasswordError:
            # If a PasswordError is raised, the PDF is protected
            return jsonify({"protected": True})
        except Exception as e:
            # Catch any other error while opening the PDF
            return jsonify({"error": f"Error checking PDF: {e}"}), 500

    except Exception as e:
        return jsonify({"error": f"Error processing PDF: {e}"}), 500


@app.route('/unlock_pdf', methods=['POST'])
def unlock_pdf():
    pdf_file = request.files.get('file')
    password = request.form.get('password')  # User-provided password
    if not pdf_file:
        return jsonify({"error": "No PDF file provided"}), 400
    if not password:
        return jsonify({"error": "No password provided"}), 400

    try:
        # Load the protected PDF with the provided password
        pdf_stream = BytesIO(pdf_file.read())
        with pikepdf.open(pdf_stream, password=password) as pdf:
            # Save the unlocked PDF to a BytesIO stream
            unlocked_pdf = BytesIO()
            pdf.save(unlocked_pdf)

        # Reset the stream pointer
        unlocked_pdf.seek(0)
        return send_file(unlocked_pdf, as_attachment=True, download_name="unlocked.pdf", mimetype="application/pdf")

    except pikepdf.PasswordError:
        return jsonify({"error": "Invalid password provided for the PDF."}), 400
    except Exception as e:
        return jsonify({"error": f"Error unlocking PDF: {e}"}), 500

# ------------------------- PDF to Word -------------------------
@app.route('/file_selection/pdf_to_word', methods=['GET'],endpoint='file_selection_pdf_to_word')
def file_selection_pdf_to_word():
    return render_template('file-selection.html',tool='pdf_to_word')

@app.route('/convert_to_word', methods=['POST'])
def convert_to_word():
    pdf_file = request.files.get('file')
    if not pdf_file:
        return "No PDF file uploaded", 400

    try:
        # Load the PDF using PyMuPDF
        pdf_stream = BytesIO(pdf_file.read())
        pdf_document = fitz.open(stream=pdf_stream, filetype="pdf")

        # Create a new Word document
        doc = Document()

        # Extract text from each page and add it to the Word document
        for page_number in range(len(pdf_document)):
            page = pdf_document[page_number]
            text = page.get_text()

            # Add a heading for each page (optional)
            doc.add_heading(f'Page {page_number + 1}', level=1)
            doc.add_paragraph(text)

        # Save the Word document to an output stream
        output_stream = BytesIO()
        doc.save(output_stream)
        pdf_document.close()

        # Reset the pointer and send the Word document as a downloadable file
        output_stream.seek(0)
        return send_file(output_stream, as_attachment=True, download_name="converted.docx", mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    except Exception as e:
        return f"Error processing PDF: {e}", 500

# ------------------------- Word to PDF -------------------------

def close_word_instances():
    """Ensure any lingering Word processes are closed."""
    for proc in psutil.process_iter(['pid', 'name']):
        if 'winword.exe' in proc.info['name'].lower():
            try:
                proc.terminate()  # Force terminate any word processes
                print(f"Terminated Word process with PID {proc.info['pid']}")
            except (psutil.NoSuchProcess, psutil.AccessDenied):
                pass  # Ignore if process is already gone or access denied

@app.route('/file_selection/word_to_pdf', methods=['GET'],endpoint='file_selection_word_to_pdf')
def file_selection_word_to_pdf():
    return render_template('file-selection.html',tool='word_to_pdf')

@app.route('/convert_word_to_pdf', methods=['POST'])
def convert_word_to_pdf():
    word_file = request.files.get('file')
    if not word_file:
        return "No Word file uploaded", 400

    try:
        # Step 1: Close any running Word processes to avoid conflicts
        close_word_instances()

        # Step 2: Save the uploaded Word file temporarily
        temp_file_path = os.path.join("uploads", word_file.filename)
        word_file.save(temp_file_path)

        # Step 3: Initialize COM explicitly before using docx2pdf
        pythoncom.CoInitialize()

        # Step 4: Convert Word to PDF using docx2pdf
        output_pdf_path = os.path.join("uploads", f"{os.path.splitext(word_file.filename)[0]}.pdf")
        convert(temp_file_path, output_pdf_path)

        # Step 5: Cleanup the temporary Word file after conversion
        os.remove(temp_file_path)

        # Step 6: Uninitialize COM to free resources
        pythoncom.CoUninitialize()

        # Step 7: Return the PDF file as a downloadable response
        return send_file(output_pdf_path, as_attachment=True, download_name=f"{os.path.splitext(word_file.filename)[0]}.pdf", mimetype="application/pdf")

    except Exception as e:
        # Ensure COM is uninitialized even if an error occurs
        pythoncom.CoUninitialize()
        return f"Error processing Word document: {e}", 500


# ------------------------- Remove Pages -------------------------
@app.route('/file_selection/remove_pages', methods=['GET'],endpoint='file_selection_remove_pages')
def file_selection_remove_pages():
    return render_template('file-selection.html',tool='remove_pages')

@app.route('/remove_pages', methods=['POST'])
def remove_pages():
    pdf_file = request.files.get('file')
    pages_to_remove = request.form.get('pages')  # Comma-separated list of pages
    if not pdf_file or not pages_to_remove:
        return "No PDF file or pages specified for removal", 400

    try:
        # Convert the comma-separated pages to a list of integers
        pages_to_remove = [int(page.strip()) - 1 for page in pages_to_remove.split(',')]

        # Read the PDF and create a new one without the specified pages
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        pdf_writer = PyPDF2.PdfWriter()

        for i in range(len(pdf_reader.pages)):
            if i not in pages_to_remove:
                pdf_writer.add_page(pdf_reader.pages[i])

        # Save the modified PDF to a BytesIO stream
        output_stream = BytesIO()
        pdf_writer.write(output_stream)

        # Reset the pointer
        output_stream.seek(0)
        return send_file(output_stream, as_attachment=True, download_name="modified.pdf", mimetype="application/pdf")

    except Exception as e:
        return f"Error processing PDF: {e}", 500

# ------------------------- Watermark PDF -------------------------
@app.route('/file_selection/add_watermark', methods=['GET'],endpoint='file_selection_add_watermark')
def file_selection():
    return render_template('file-selection.html',tool='add_watermark')

# Function to create a watermark PDF (for text or image)
def create_watermark(watermark_text=None, watermark_image=None, position="center", page_size=(595.27, 841.89)):
    packet = BytesIO()
    can = canvas.Canvas(packet, pagesize=page_size)

    # Positions for the watermark
    positions = {
        "top-left": (50, page_size[1] - 50),  # Top-left corner
        "top-right": (page_size[0] - 150, page_size[1] - 50),
        "bottom-left": (50, 50),
        "bottom-right": (page_size[0] - 150, 50),
        "center": (page_size[0] / 2 - 75, page_size[1] / 2),
    }

    # Get coordinates for the selected position
    x, y = positions.get(position, (page_size[0] / 2, page_size[1] / 2))

    # Add watermark text if provided
    if watermark_text:
        can.setFont("Helvetica-Bold", 36)
        can.setFillGray(0.5, 0.5)  # Light gray for better readability
        can.drawString(x, y, watermark_text)

    # Add watermark image if provided
    if watermark_image:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as temp_image_file:
            temp_image_file.write(watermark_image.read())
            temp_image_path = temp_image_file.name

        # Draw the image with transparency and ensure it doesn't fully obscure text
        can.saveState()
        can.setFillAlpha(0.2)  # Set transparency (0.2 is semi-transparent)
        can.drawImage(temp_image_path, x, y, width=200, height=200, mask='auto')  # Adjust width and height
        can.restoreState()

        # Remove the temporary image file
        os.remove(temp_image_path)

    can.save()
    packet.seek(0)
    return packet

# Function to overlay watermark onto an existing PDF
def add_watermark(input_pdf: BytesIO, watermark_pdf: BytesIO):
    pdf_reader = PdfReader(input_pdf)
    pdf_writer = PdfWriter()

    watermark_reader = PdfReader(watermark_pdf)
    watermark_page = watermark_reader.pages[0]

    for page in pdf_reader.pages:
        page.merge_page(watermark_page)
        pdf_writer.add_page(page)

    output_pdf = BytesIO()
    pdf_writer.write(output_pdf)
    output_pdf.seek(0)
    return output_pdf

@app.route('/apply_watermark', methods=['POST'])
def apply_watermark():
    pdf_file = request.files.get('file')
    watermark_text = request.form.get('watermark_text')
    watermark_image = request.files.get('watermark_image')
    position = request.form.get('position', 'center')

    if not pdf_file:
        return "No PDF file uploaded", 400

    try:
        # Read the uploaded PDF file
        pdf_stream = BytesIO(pdf_file.read())

        # Create the watermark
        if watermark_text or watermark_image:
            watermark_pdf_stream = create_watermark(
                watermark_text=watermark_text.strip() if watermark_text else None,
                watermark_image=BytesIO(watermark_image.read()) if watermark_image else None,
                position=position,
                page_size=(595.27, 841.89)  # Default A4 page size
            )
        else:
            return "No watermark text or image provided", 400

        # Apply the watermark to the PDF
        output_pdf = add_watermark(pdf_stream, watermark_pdf_stream)

        # Send the updated PDF back to the user
        return send_file(output_pdf, as_attachment=True, download_name="watermarked.pdf", mimetype="application/pdf")

    except Exception as e:
        return f"Error applying watermark: {e}", 500

if __name__ == '__main__':
    app.run(debug=True)
